# -*- coding: utf-8 -*-
"""生成《业务外包临时用工人员申请表》Word 版（签字用首页）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, name="仿宋_GB2312", size=16, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=16, bold=False,
                  first_line_indent=None, space_before=0, space_after=0, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def main():
    out = Path(__file__).resolve().parents[1] / "业务外包临时用工人员申请表_4人_签字用.docx"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "仿宋_GB2312"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    style.font.size = Pt(16)

    add_paragraph(doc, "附件 2", size=14, space_after=6)

    title = add_paragraph(
        doc,
        "业务外包临时用工人员申请表",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=22,
        bold=True,
        space_before=12,
        space_after=18,
        line_spacing=1.2,
    )
    set_run_font(title.runs[0], name="小标宋", size=22, bold=True)

    body = (
        "    因现场生产作业需要，需增加业务外包工 4 人，具体明细见"
        "《兰州西工务段业务外包工临时用工人员信息库》（其中马连连此前已申报备案），"
        "经职教科对业务外包工培训资料进行审核，车间、主管科室、安全科对业务外包工"
        "培训情况、保险购买、年龄等资料进行审核，业务外包工人员资料齐全有效，"
        "同意业务外包工使用。"
    )
    add_paragraph(doc, body, first_line_indent=0.74, space_after=18, line_spacing=1.75)

    sign_lines = [
        "车间主任、书记审核：",
        "职教科审核：",
        "主管科室审核：",
        "安全科审核：",
        "安全副段长审批：",
    ]
    for line in sign_lines:
        p = add_paragraph(doc, f"    {line}", space_after=16, line_spacing=1.5)
        p.add_run(" " * 18)

    footer_block = doc.add_paragraph()
    footer_block.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_block.paragraph_format.space_before = Pt(24)
    footer_block.paragraph_format.space_after = Pt(18)
    footer_block.paragraph_format.line_spacing = 1.5
    r1 = footer_block.add_run("白银桥隧车间\n")
    set_run_font(r1, size=16)
    r2 = footer_block.add_run("2026 年 08 月 26 日")
    set_run_font(r2, size=16)

    note = (
        "    说明：业务外包临时用工人员申请表逐级审批同意后，由安全科给予相关"
        "人员在工务安全生产管理系统及营业线施工作业监控系统维护。"
    )
    add_paragraph(doc, note, size=14, first_line_indent=0.74, space_before=6, line_spacing=1.5)

    page_no = add_paragraph(doc, "- 1 -", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_before=18)

    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
